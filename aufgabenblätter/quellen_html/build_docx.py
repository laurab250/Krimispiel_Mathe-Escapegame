# -*- coding: utf-8 -*-
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image, ImageDraw

OUT="/home/tim/Downloads/tst/aufgabenblätter"
GRID="/tmp/grids"; os.makedirs(GRID, exist_ok=True)

def make_grid(h_mm):
    path=f"{GRID}/g{h_mm}.png"
    if os.path.exists(path): return path
    cells_x=32; cell=24; rows=max(1,round(h_mm/5))
    W=cells_x*cell; H=rows*cell
    img=Image.new("RGB",(W+1,H+1),"white"); dr=ImageDraw.Draw(img); col=(150,180,215)
    for x in range(0,W+1,cell): dr.line([(x,0),(x,H)],fill=col,width=1)
    for y in range(0,H+1,cell): dr.line([(0,y),(W,y)],fill=col,width=1)
    img.save(path,dpi=(96,96)); return path

def shade(cell,hexfill):
    sh=OxmlElement('w:shd'); sh.set(qn('w:val'),'clear'); sh.set(qn('w:fill'),hexfill)
    cell._tc.get_or_add_tcPr().append(sh)

def add_runs(p, text):
    # parse **bold**
    for i,part in enumerate(text.split("**")):
        if part=="": continue
        r=p.add_run(part); r.bold=(i%2==1)

def box(doc, label, text, fill="EEF1F4"):
    t=doc.add_table(rows=1,cols=1); t.alignment=WD_TABLE_ALIGNMENT.CENTER
    t.style='Table Grid'; c=t.cell(0,0); shade(c,fill)
    p=c.paragraphs[0]
    if label:
        r=p.add_run(label+"  "); r.bold=True; r.font.size=Pt(8.5)
    add_runs(p,text)
    for r in p.runs:
        if not r.font.size: r.font.size=Pt(10)
    doc.add_paragraph().paragraph_format.space_after=Pt(2)

def h1(doc,title,sub):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(0)
    r=p.add_run(title); r.bold=True; r.font.size=Pt(16)
    p2=doc.add_paragraph(); r=p2.add_run(sub+"  ·  Fall #2024-0311"); r.italic=True; r.font.size=Pt(9); r.font.color.rgb=RGBColor(0x55,0x55,0x55)
    pb=p2._p.get_or_add_pPr(); bd=OxmlElement('w:pBdr'); b=OxmlElement('w:bottom')
    b.set(qn('w:val'),'single'); b.set(qn('w:sz'),'12'); b.set(qn('w:space'),'2'); b.set(qn('w:color'),'000000'); bd.append(b); pb.append(bd)
    m=doc.add_paragraph(); r=m.add_run("Name(n): ________________________________     Klasse: __________     Datum: __________"); r.font.size=Pt(9)
    m.paragraph_format.space_after=Pt(4)

def h2(doc,txt):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(6); p.paragraph_format.space_after=Pt(1)
    r=p.add_run(txt); r.bold=True; r.font.size=Pt(11.5)

def para(doc,txt):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(2); add_runs(p,txt)
    for r in p.runs:
        if not r.font.size: r.font.size=Pt(10.5)

def fields(doc, items):
    for lab in items:
        p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(1)
        r=p.add_run(lab+" "); r.font.size=Pt(10.5)
        r2=p.add_run("_"*max(6,int((85-len(lab))/1.6))); r2.font.size=Pt(10.5)

def lines(doc,n):
    for _ in range(n):
        p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(6)
        pb=p._p.get_or_add_pPr(); bd=OxmlElement('w:pBdr'); b=OxmlElement('w:bottom')
        b.set(qn('w:val'),'single'); b.set(qn('w:sz'),'6'); b.set(qn('w:space'),'1'); b.set(qn('w:color'),'999999'); bd.append(b); pb.append(bd)

def grid(doc,h):
    doc.add_picture(make_grid(h), width=Cm(16))
    doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER

def table(doc, headers, rows, widths=None):
    t=doc.add_table(rows=1,cols=len(headers)); t.style='Table Grid'
    hdr=t.rows[0].cells
    for i,htext in enumerate(headers):
        shade(hdr[i],"E3E3E3"); p=hdr[i].paragraphs[0]; r=p.add_run(htext); r.bold=True; r.font.size=Pt(9.5)
    for row in rows:
        cells=t.add_row().cells
        for i,val in enumerate(row):
            p=cells[i].paragraphs[0]; r=p.add_run(val); r.font.size=Pt(9.5)
            cells[i].paragraphs[0].paragraph_format.space_after=Pt(3)
    doc.add_paragraph().paragraph_format.space_after=Pt(0)

def new(title,sub):
    d=Document()
    s=d.sections[0]; s.top_margin=Cm(1.3); s.bottom_margin=Cm(1.1); s.left_margin=Cm(2.2); s.right_margin=Cm(2.2)
    st=d.styles['Normal']; st.font.name='Calibri'; st.font.size=Pt(10.5)
    h1(d,title,sub); return d

S={}
# ---- 1 ----
d=new("SPUR 1 — TATORTANALYSE","Polizeibericht · Überwachungsbild")
box(d,"AUFTRAG","Ihr habt Zugang zum Überwachungsbild des Labors. Beantwortet die Fragen so genau wie möglich. Haltet euch **ausschließlich an das, was ihr im Bild seht**. Diese Beobachtungen braucht ihr in späteren Hinweisen wieder.")
h2(d,"Beobachtungen")
fields(d,["1. Uhrzeit (Wanduhr):","2. Gesicherte Tür – Aufschrift & Bedeutung:"])
para(d,"3. Drei Zugangschips auf Dr. Felds Schreibtisch – welche Nummer deutet auf wen?")
fields(d,["Chip #____ →","Chip #____ →","Chip #____ →","Welcher Chip liegt nicht am richtigen Ort?"])
fields(d,["4. Notizbuch – Text & mögliche Person:","5. Geplantes Gespräch mit Dr. Feld – Person & Uhrzeit:",
"6. Kaffeetasse mit Kürzel – Kürzel/Person/warum auffällig:","7. Handgeschriebene Notiz – Inhalt:",
"8. „TTX“ an zwei Stellen – Person/Kontext:","9. Wartungsschild am Sicherheitspanel – Name:",
"10. Vorlesungs-Alibi – welcher Hinweis bestätigt es?"])
h2(d,"Technische Daten (für Spur 5)")
fields(d,["11. Kamera-Typenschild – Montagehöhe:","      Neigungswinkel:","      Standort laut Schild:"])
h2(d,"Erste Einschätzung")
para(d,"Welche zwei Personen erscheinen euch am verdächtigsten? Jeweils mit Begründung aus dem Bild.")
lines(d,4); S["1_Polizeibericht"]=d
# ---- 2 ----
d=new("SPUR 2 — HANDGRÖSSEN","Normalverteilung · Handschuhfund")
box(d,"GEGEBEN","Vergleichsstudie: **μ = 19,2 cm**, **σ = 1,5 cm**. Die gesicherten Handschuhe passen zu einer Handgröße von **18,5 cm bis 21,5 cm**.")
h2(d,"a) Standardisierung – z-Wert z = (x − μ) / σ")
table(d,["Person","Handgröße x [cm]","z-Wert"],
 [["Jonas Kern","19,8",""],["Miriam Scholz","19,3",""],["Dr. Elena Voss","18,9",""],["Prof. Anton Richter","20,5",""],
  ["Sarah Mayer","18,7",""],["Felix Hauser","20,4",""],["Lena Brandt","18,6",""],["David Okonkwo","22,0",""]])
grid(d,45)
h2(d,"b) Auswertung")
para(d,"Wer liegt **außerhalb** des Handschuh-Bereichs (18,5–21,5 cm)? Wer hat den größten Betrag von z?")
grid(d,35)
h2(d,"c) Schlussfolgerung")
box(d,"","Welche Person kann den Handschuh nicht getragen haben und scheidet aus? Begründung:")
lines(d,3); S["2_Handgroessen"]=d
# ---- 3a ----
d=new("SPUR 3a — TODESZEITPUNKT","Abkühlungsgesetz · Exponentialfunktion")
box(d,"GEGEBEN","**T(t) = 20 + 17 · e^(−0,16·t)**  (T in °C, t in Stunden seit dem Tod). Messung um **18:00 Uhr: T = 30,0 °C**, Raumtemperatur 20 °C.")
h2(d,"Aufgabe")
para(d,"Bestimmt den Todeszeitpunkt: Löst T(t) = 30 nach t (natürlicher Logarithmus) und rechnet in eine Uhrzeit um.")
grid(d,95)
box(d,"ERGEBNIS","t = ____________ Stunden vor 18:00 Uhr   ⟹   Todeszeitpunkt ≈ ____________ Uhr")
para(d,"Dieses Ergebnis legt das Tatzeitfenster fest – ihr braucht es in Spur 6 und 7 wieder.")
S["3a_Todeszeit"]=d
# ---- 3b ----
d=new("SPUR 3b — HERZFREQUENZ","Änderungsrate · Differenzialrechnung")
box(d,"GEGEBEN","Herzfrequenz **f(t)** in Schlägen/Minute, t in Stunden ab 13:00 Uhr. Tatzeitpunkt **t = 1,7** (14:42 Uhr). Wer aktiv/nervös war, zeigt eine deutliche Rate; wer ruhig war, eine Rate nahe 0.")
h2(d,"a) Ableitungen – f′(1,7)")
table(d,["Person","f(t)","f′(t)","f′(1,7)"],
 [["Jonas Kern","2,5t⁴ − 15t³ + 25t² + 70","",""],
  ["Dr. Elena Voss","5,15t⁴ − 29,52t³ + 46,35t² + 72","",""],
  ["Prof. A. Richter","13,35t⁴ − 73,0t³ + 106,8t² + 68","",""],
  ["Sarah Mayer","7,475t⁴ − 43,86t³ + 71,01t² + 65","",""],
  ["Felix Hauser","0,3t³ − 1,2t² + 1,2t + 71,7","",""]])
grid(d,55)
h2(d,"b) Sonderfall Miriam Scholz")
para(d,"Nur vier Werte ihrer Funktion **g(t) = a·t³ + b·t² + c·t + d** sind gesichert: g(0) = 65 · g(2) = 85 · g′(2) = 0 · g″(2) = −8. Stellt das Gleichungssystem auf, bestimmt g(t) und berechnet g′(1,7).")
grid(d,85)
h2(d,"c) Schlussfolgerung")
box(d,"","Wessen Herzfrequenz zeigte zur Tatzeit praktisch keine Änderung (Rate ≈ 0)? Diese Person scheidet aus:")
lines(d,2); S["3b_Herzfrequenz"]=d
# ---- 4 ----
d=new("SPUR 4 — GPS-AUSWERTUNG","Bestimmtes Integral · Alibiüberprüfung")
box(d,"GEGEBEN","Geschwindigkeitsfunktionen **v(t)** in m/min, t in Minuten ab 14:00 Uhr. Das Integral von v(t) ergibt die zurückgelegte **Strecke**. Die Distanz laut Lageplan entspricht dem behaupteten Aufenthaltsort.")
h2(d,"a) Strecke = ∫ v(t) dt")
table(d,["Person","v(t) [m/min]","t-Bereich","Distanz (Lageplan)","Strecke [m]"],
 [["Jonas Kern","−0,016t² + 1,2t − 20","[25, 50]","80 m",""],
  ["Miriam Scholz","−0,025t² + 1,875t − 27,5","[20, 55]","380 m",""],
  ["Dr. E. Voss","−0,05t² + 2,95t − 36,9","[18, 41]","220 m",""],
  ["Prof. Richter","−0,095t² + 2,85t − 11,875","[5, 25]","600 m",""],
  ["Lena Brandt","−0,04t² + 0,6t","[0, 15]","120 m",""],
  ["Sarah Mayer","−0,1t² + 9t − 180","[30, 60]","450 m",""]])
grid(d,80)
h2(d,"b) Vergleich & Schlussfolgerung")
para(d,"Bei wem stimmt die berechnete Strecke mit der Lageplan-Distanz überein (Alibi bestätigt)?")
box(d,"","Wessen Aussage wird durch die GPS-Daten bestätigt und scheidet aus? Bei wem passt die Strecke nicht?")
lines(d,2); S["4_GPS"]=d
# ---- 5 ----
d=new("SPUR 5 — KÖRPERGRÖSSE","Trigonometrie · Kameraauswertung")
box(d,"GEGEBEN","Südwand-Kamera (Daten aus Spur 1): Montagehöhe **h = 3,20 m**, Neigungswinkel **α = 16°**, horizontaler Abstand **d = 5,0 m**. Die Kamera erfasst den Kopf der Person. Messungenauigkeit **± 10 cm**.")
h2(d,"a) Körpergröße berechnen")
para(d,"Skizze (Kamera, Sichtlinie, Person) und Rechnung. Hinweis: Die Sichtlinie fällt über d um d·tan(α) ab.")
grid(d,80)
h2(d,"b) Abgleich (Toleranz ± 10 cm)")
table(d,["Person","Körpergröße [cm]","im Bereich?"],
 [["Jonas Kern","178",""],["Dr. Elena Voss","169",""],["Prof. Anton Richter","181",""],["Lena Brandt","162",""],["Miriam Scholz","174",""]])
h2(d,"c) Schlussfolgerung")
box(d,"","Welche Person liegt außerhalb des Toleranzbereichs und scheidet aus? Begründung:")
lines(d,2); S["5_Koerpergroesse"]=d
# ---- 6 ----
d=new("SPUR 6 — BLUTZUCKER","Exponentialfunktion · Glukosesensor")
box(d,"GEGEBEN","**B(t) = 95 · e^(−0,03·t) + 85**  (mg/dL, t in Minuten ab 14:00 Uhr). Grenzwert **135 mg/dL** (darüber = kürzlich gegessen). Küche 11:30–14:30 Uhr. Tatzeitfenster 14:28–14:47 Uhr.")
h2(d,"a) Funktionswerte")
fields(d,["B(0) = __________  Bedeutung:","B(30) = __________  Bedeutung:"])
grid(d,35)
h2(d,"b) Grenzwert")
para(d,"Ab wann liegt der Blutzucker unter 135 mg/dL? Gleichung aufstellen und nach t lösen (ln).")
grid(d,70)
box(d,"ERGEBNIS","t = ____________ Minuten ab 14:00 Uhr = ____________ Uhr")
h2(d,"c) Schlussfolgerung")
para(d,"Kombiniert mit Küchen-Öffnung, GPS (Spur 4) und Tatzeitfenster: Was folgt für Richters Aufenthaltsort – und warum scheidet er aus?")
lines(d,3); S["6_Blutzucker"]=d
# ---- 7 ----
d=new("SPUR 7 — ZUGANGSSYSTEM","Lineares Gleichungssystem · Bereich B")
box(d,"AUFTRAG","Hört euch den Anruf der Polizei (KHK Brenner) aufmerksam an. Alle Angaben – Bedeutung von p und t, die zwei Formeln und die Prüfwerte – entnehmt ihr dem Gespräch.")
h2(d,"1) Mitschrift während des Anrufs")
fields(d,["p (Chip-Nummer) bedeutet:","t (Zeitwert) bedeutet:","Formel 1:","Formel 2:","Prüfwerte: C1 = __________   C2 ="])
h2(d,"2) Wer hatte Zugang zu Bereich B?")
table(d,["Chip-Nr.","Person","Sonderausweis?","noch verdächtig?"],
 [["#1","Dr. Markus Feld (Opfer)","Ja","Tatopfer – entfällt"],["#3","","",""],["#7","","",""],["keiner","Dr. Elena Voss","",""]])
fields(d,["Wer scheidet allein durch den fehlenden Sonderausweis aus?"])
h2(d,"3) Gleichungssystem aufstellen")
fields(d,["(I)","(II)"])
h2(d,"4) Gleichungssystem lösen")
grid(d,90)
h2(d,"5) Ergebnis deuten")
box(d,"","p = ______ ⟹ Chip #______ gehört ____________________      t = ______ ⟹ 14:00 + ______ min = ______ Uhr.   Im Tatzeitfenster (14:28–14:47)? ______")
h2(d,"6) Täter benennen und begründen")
para(d,"Täter und mindestens **drei** mathematische Argumente aus verschiedenen Hinweisen.")
lines(d,5); S["7_Zugangssystem"]=d

for n,doc in S.items():
    doc.save(os.path.join(OUT,n+".docx"))
print("DOCX erstellt:", ", ".join(sorted(S)))
